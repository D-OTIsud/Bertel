# -*- coding: utf-8 -*-
"""Build the Bertel RGPD/DPO documentation pack.

The Markdown sources and DOCX deliverables are generated together so the pack can
be regenerated after DPO/legal review without hand-editing Word files.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "livrables"
DATE = "13 juin 2026"


OFFICIAL_SOURCES = """\
## Sources officielles consultées

- CNIL - DPO : par où commencer ? https://www.cnil.fr/fr/dpo-par-ou-commencer
- CNIL - Le registre des activités de traitement : https://cnil.fr/fr/RGPD-le-registre-des-activites-de-traitement
- CNIL - L'analyse d'impact relative à la protection des données (AIPD) : https://www.cnil.fr/fr/RGPD-analyse-impact-protection-des-donnees-aipd
- CNIL - Outil PIA : https://www.cnil.fr/fr/outil-pia-telechargez-et-installez-le-logiciel-de-la-cnil
- CNIL - Les droits des personnes sur leurs données : https://www.cnil.fr/fr/passer-laction/les-droits-des-personnes-sur-leurs-donnees
- CNIL - Violations de données personnelles : les règles à suivre : https://www.cnil.fr/fr/violations-de-donnees-personnelles-les-regles-suivre
- CNIL - Guide pratique RGPD pour les DPO : https://www.cnil.fr/sites/default/files/atoms/files/guide_pratique_rgpd_-_delegues_a_la_protection_des_donnees.pdf
"""


PROJECT_SOURCES = """\
## Sources projet utilisées

- `README.md` : positionnement Bertel, stack Next.js + Supabase/PostgreSQL, API et documentation.
- `ARCHITECTURE.md` et `docs/architecture/bertel-object-workspace-canonical-map.md` : carte fonctionnelle du workspace objet.
- `docs/architecture/OBJECT_DATA_DICTIONARY.md` : dictionnaire des données objet, contacts, médias, CRM, incidents, juridique, RLS et audit.
- `db-graph-out/DB_AGENT_INDEX.md`, `FUNCTIONS.md`, `POLICIES.md`, `TYPES.md` : cartographie DB, RPC, politiques RLS et enums.
- `dbdoc/*.md` : fiches de tables sensibles (`auth.users`, `app_user_profile`, `actor`, `actor_channel`, `actor_consent`, `contact_channel`, `object_private_description`, `crm_interaction`, `crm_task`, `incident_report`, `media`, `object_legal`, `audit.audit_log`).
- `bertel-tourism-ui/src/views/*` et `bertel-tourism-ui/src/services/*` : parcours UI, authentification, RBAC, CRM, modération, audits, médias.
"""


BADGEUSE_SOURCES = """\
## Sources internes transposées depuis le dossier badgeuse

- `C:\\Users\\dphil\\Downloads\\Manuel_Utilisateur_Badgeuse_OTI (7).pdf` : modèle de manuel utilisateur interne OTI.
- `C:\\Users\\dphil\\Downloads\\Reglement_RGPD_Badgeuse_OTISUD (3).pdf` : identité SPL OTI DU SUD, référent RGPD, hébergement OVH France, absence de transfert hors UE et mesures de sécurité.
- `C:\\Users\\dphil\\Downloads\\DPIA_Badgeuse_OTISUD (3).pdf` : modèle d'analyse proportionnée, risque résiduel faible, Supabase/PostgreSQL, HTTPS, RLS et logs administratifs.
"""


ORG_CONTEXT = """\
## Informations organisationnelles reprises du dossier badgeuse

| Élément | Valeur |
|---|---|
| Organisation | SPL OTI DU SUD |
| Adresse | 379 Rue Hubert Delisle, 97430 Le Tampon |
| SIREN | 882 699 556 |
| Activité | Tourisme |
| Effectif de référence | 27 salariés |
| Date de création | 24/01/2020 |
| Référent RGPD | David Philippe - Manager SI |
| Contact RGPD | d.philippe@otisud.com - 06 93 41 92 91 |
| Hébergement de référence | OVH France, UE, sans transfert hors UE, à confirmer si l'architecture Bertel diffère |
"""


PUBLIC_DATA_POSTURE = """\
## Posture données Bertel

En l'état du périmètre décrit, environ 95 % des informations stockées dans Bertel sont publiques ou destinées à être publiées : fiches touristiques, descriptions, horaires, coordonnées professionnelles d'établissement, médias de promotion, classifications, tarifs, équipements et statuts de publication. Le risque RGPD principal ne vient donc pas du référentiel public lui-même, mais des zones internes : interactions entre l'office de tourisme et les prestataires, notes CRM, tâches, historiques de relation, coordonnées personnelles lorsqu'elles diffèrent des coordonnées publiques ou professionnelles, et erreurs de publication accidentelles.
"""


@dataclass(frozen=True)
class PackDoc:
    filename: str
    title: str
    subtitle: str
    body: str

    @property
    def md_name(self) -> str:
        return f"{self.filename}.md"

    @property
    def docx_name(self) -> str:
        return f"{self.filename}.docx"


DOCUMENTS: list[PackDoc] = [
    PackDoc(
        filename="01_Manuel_utilisateur_Bertel",
        title="Manuel utilisateur Bertel",
        subtitle="Parcours métier, rôles, bonnes pratiques RGPD et exploitation quotidienne",
        body=f"""\
> Statut : document de travail à valider par l'équipe produit, le responsable de traitement et le DPO.
> Version : {DATE}. Périmètre : Bertel 3.0, UI Next.js, backend Supabase/PostgreSQL.

## 1. Objet du manuel

Bertel est la Base d'Enregistrement et de Référentiel Touristique des Établissements et Lieux. L'application centralise des fiches touristiques, des médias, des contacts, des relations avec les prestataires, des workflows de modération, des audits, des signalements et des informations de conformité.

Ce manuel s'adresse aux agents de l'office de tourisme, administrateurs d'organisation, administrateurs plateforme, contributeurs autorisés, référents qualité et DPO. Il décrit les actions attendues dans l'interface et les règles minimales à appliquer pour ne pas exposer de données personnelles inutiles.

{PUBLIC_DATA_POSTURE}

## 2. Connexion et session

1. Accéder à l'URL de l'application Bertel.
2. Se connecter avec Google OAuth lorsque le fournisseur est actif, ou avec email et mot de passe lorsque ce mode est configuré.
3. Vérifier que le nom, l'organisation active et les droits affichés correspondent au contexte de travail.
4. Se déconnecter après usage sur un poste partagé.

Le compte applicatif est relié à `auth.users` et `app_user_profile`. Le profil contient notamment l'email d'authentification, le nom d'affichage, l'avatar, la langue, le fuseau horaire, le rôle plateforme et les préférences.

## 3. Rôles et accès

| Profil | Usage principal | Points de vigilance |
|---|---|---|
| `owner` | Gouvernance plateforme et actions d'administration globale. | Accès très large ; réservé aux personnes habilitées. |
| `super_admin` | Administration technique/fonctionnelle transverse. | Peut accéder à des vues sensibles ; utiliser un compte nominatif. |
| `tourism_agent` | Exploitation métier quotidienne, selon permissions et organisation active. | Les droits effectifs dépendent de l'organisation, des rôles métier/admin et des permissions. |
| Admin d'organisation | Administration de l'équipe et des permissions d'une ORG. | Ne pas accorder plus de droits que nécessaire. |

La page Équipe permet de gérer les membres, les rôles métier, les rôles admin et les permissions. Les rôles d'administration utilisent un rang ; un administrateur ne doit pas modifier son propre rôle ou ses propres permissions.

## 4. Navigation principale

| Zone | Finalité | Données manipulées |
|---|---|---|
| Dashboard | Vue de pilotage qualité, offre, actualisation et filtres. | Agrégats, indicateurs, données de fiches touristiques. |
| Explorer | Recherche, carte, filtres et consultation des fiches. | Objets publiés, brouillons autorisés, médias, localisation, badges. |
| Fiche objet / édition | Création et modification des établissements, lieux, activités, événements, itinéraires et services. | Informations publiques, internes, légales et relationnelles. |
| Modération | Comparaison avant/après et décision sur les changements proposés. | `pending_change`, auteur, notes de revue, valeurs modifiées. |
| CRM | Annuaire acteurs, tâches, relances et timeline relationnelle. | Acteurs, contacts, interactions, tâches, notes métier. |
| Audits | Checklists terrain et incidents. | Résultats d'audit, signalements, photos, géolocalisation, déclarants. |
| Publications | Sélection et suivi des objets destinés aux supports édités. | Statuts de publication, rattachements. |
| Paramètres | Marque blanche, thèmes, paramètres d'application. | Paramètres visuels et d'organisation. |

## 5. Rechercher et consulter une fiche

1. Ouvrir Explorer.
2. Utiliser les filtres par type, commune, accessibilité, durabilité, horaires, labels ou carte.
3. Sélectionner une fiche dans la liste ou sur la carte.
4. Vérifier le statut : publié, brouillon visible aux éditeurs, archivé ou masqué selon droits.
5. Ne pas copier des informations internes dans des supports publics sans vérifier leur visibilité.

Les objets `ORG` sont des entités internes de rattachement institutionnel et ne sont pas exposés comme un type de recherche grand public.

## 6. Modifier une fiche objet

Les onglets suivent une logique métier : identité, taxonomie, localisation, descriptions, médias, contacts publics, caractéristiques, tarifs, horaires, relation prestataires, juridique, conformité, audits et modules conditionnels selon type.

Bonnes pratiques :

- Enregistrer par onglet lorsque le travail est terminé.
- Lire les messages de validation avant de quitter un onglet.
- Utiliser la modération lorsqu'un changement doit être revu avant publication.
- Distinguer les champs publics des champs internes.
- Vérifier la langue de chaque champ traduisible.
- Ne pas stocker dans une note interne des données sensibles ou hors sujet.

## 7. Médias et droits

Chaque média peut être rattaché à l'objet global ou à un sous-lieu. Il contient un type, un titre, une description, un crédit, une URL, une visibilité, un indicateur de publication, un média principal et une date d'expiration des droits.

Avant de publier un média :

- vérifier que les droits d'usage sont documentés ;
- renseigner le crédit lorsque nécessaire ;
- éviter les visages reconnaissables, plaques, documents administratifs ou informations personnelles visibles ;
- classer la visibilité : public, partenaires ou interne ;
- ajouter une description utile à l'accessibilité si le média est diffusé.

## 8. Contacts, acteurs et consentements

Bertel distingue :

- les contacts publics d'un établissement (`contact_channel`) : téléphone, email, site web, réseaux sociaux publiables ;
- les acteurs (`actor`) et leurs canaux (`actor_channel`) : personnes ou contacts liés à un objet ou une organisation ;
- les consentements (`actor_consent`) : préférences et preuves utiles à la relation.

Règles d'usage :

- publier uniquement les coordonnées destinées au public ;
- conserver les coordonnées personnelles dans les zones internes prévues ;
- documenter le consentement lorsque le traitement repose sur lui ;
- supprimer ou archiver les contacts obsolètes après validation métier ;
- ne pas utiliser les contacts Bertel pour des finalités non prévues.

## 9. CRM, notes privées et relances

Le module CRM Bertel est un suivi relationnel métier, pas un pipeline commercial. Il sert à garder une trace des échanges avec les prestataires, des demandes, des relances et des tâches.

Les notes privées sont utiles pour assurer la continuité de service, mais elles doivent rester factuelles, proportionnées et non discriminatoires. Ne pas écrire d'appréciations personnelles inutiles, d'informations de santé, d'opinions politiques, de données familiales ou de jugements subjectifs sur une personne.

## 10. Audits et incidents

Les audits documentent des contrôles qualité. Les incidents documentent un signalement terrain : catégorie, gravité, statut, description, localisation, déclarant et médias éventuels.

À la saisie d'un incident :

1. Décrire le fait observable.
2. Indiquer la localisation seulement si elle est utile.
3. Éviter les données personnelles dans les photos et descriptions.
4. Ne collecter le nom/email du déclarant que si un suivi est nécessaire.
5. Passer en revue les pièces jointes avant diffusion ou partage.

## 11. Administration d'équipe

La page Équipe permet :

- d'inviter un membre ;
- d'associer un utilisateur à une organisation ;
- d'attribuer un rôle métier ;
- d'attribuer ou retirer un rôle admin ;
- d'accorder ou retirer des permissions ;
- de désactiver un membre.

Chaque action d'administration doit respecter le principe de moindre privilège. Les comptes doivent être nominatifs, désactivés au départ d'une personne et revus régulièrement.

## 12. Règles RGPD au quotidien

- Minimiser : collecter seulement ce qui sert au référentiel touristique, à la relation prestataire, à la qualité ou à la conformité.
- Cloisonner : ne pas placer des informations internes dans les champs publics.
- Vérifier : contrôler les contacts, médias et notes avant publication.
- Tracer : utiliser les workflows de modération et d'audit plutôt que des échanges informels.
- Protéger : ne pas partager d'exports hors des canaux autorisés.
- Signaler : tout incident de sécurité ou exposition involontaire doit être remonté immédiatement au DPO ou au référent désigné.

## 13. Support et escalade

| Situation | Action utilisateur | Escalade |
|---|---|---|
| Accès impossible | Vérifier compte, navigateur, session, URL. | Administrateur organisation ou plateforme. |
| Donnée personnelle affichée publiquement par erreur | Dépublier ou masquer si habilité, puis signaler. | DPO + responsable métier. |
| Fiche incorrecte ou obsolète | Corriger ou créer une demande modérée. | Référent qualité si conflit. |
| Suspicion de fuite ou accès non autorisé | Ne pas exporter, capturer les éléments factuels. | DPO immédiatement, procédure violation. |
| Demande d'accès/suppression d'une personne | Ne pas répondre seul si le périmètre est incertain. | DPO ou canal officiel droits des personnes. |

{PROJECT_SOURCES}

{BADGEUSE_SOURCES}
""",
    ),
    PackDoc(
        filename="02_Reglement_interne_RGPD_DPO_Bertel",
        title="Règlement interne de conformité RGPD et gouvernance DPO",
        subtitle="Règles d'utilisation, responsabilités, contrôles et preuves de conformité pour Bertel",
        body=f"""\
> Statut : modèle interne à valider. Ce document ne remplace pas l'avis du DPO, du juriste ou du responsable de traitement.
> Version : {DATE}. Informations organisationnelles reprises du dossier badgeuse OTI SUD, à confirmer si Bertel diffère.

## 1. Périmètre

Le présent règlement s'applique à Bertel, référentiel touristique et espace de travail de gestion des fiches, médias, contacts, acteurs, relations prestataires, audits, incidents, publications, rôles, permissions et traces techniques.

Il couvre les usages internes, les exports, les partages à des partenaires, la diffusion publique et les traitements réalisés via Supabase/PostgreSQL, l'UI Next.js, les API, les imports et les outils associés.

{ORG_CONTEXT}

{PUBLIC_DATA_POSTURE}

## 2. Gouvernance

| Rôle | Responsabilité |
|---|---|
| Responsable de traitement | SPL OTI DU SUD - 379 Rue Hubert Delisle, 97430 Le Tampon - SIREN 882 699 556. |
| DPO / référent protection des données | David Philippe - Manager SI - d.philippe@otisud.com - 06 93 41 92 91. |
| Responsable produit Bertel | Maintient la cohérence métier, priorise les corrections de conformité. |
| Administrateur plateforme | Gère configuration, sécurité, accès techniques, logs et incidents. |
| Administrateur d'organisation | Gère membres, rôles et permissions dans son périmètre. |
| Utilisateur métier | Saisit, corrige et consulte les données en respectant le présent règlement. |
| Sous-traitants | OVH France, Supabase/PostgreSQL et services associés à confirmer ; le dossier badgeuse indique OVH France, sans transfert hors UE. |

Le DPO doit pouvoir être contacté facilement par les personnes concernées et par les utilisateurs internes. Le DPO tient ou supervise le registre, les AIPD, les procédures de droits et les violations.

## 3. Principes applicables

Bertel doit respecter les principes suivants :

- finalité déterminée : les données servent au référentiel touristique, à la relation prestataire, à la qualité, à la publication, à la conformité et à la sécurité ;
- minimisation : aucune donnée personnelle non nécessaire ne doit être saisie ;
- exactitude : les contacts, rôles, statuts et informations publiées doivent être mis à jour ;
- limitation de conservation : chaque famille de données doit avoir une durée validée ;
- intégrité et confidentialité : accès par rôle, RLS, audit, cloisonnement public/interne ;
- responsabilité : les choix, contrôles et incidents doivent être documentés.

## 4. Catégories de données Bertel

| Catégorie | Exemples | Sensibilité |
|---|---|---|
| Comptes utilisateurs | email, nom, avatar, rôle, langue, organisation, permissions, session. | Données personnelles internes. |
| Prestataires et acteurs | nom, fonction, contacts, rattachement à un objet, consentements. | Principalement professionnel ; sensible si coordonnées personnelles ou consentements. |
| Contacts publics | téléphone, email, site web, réseaux sociaux d'un établissement. | Donnée professionnelle publiable si destinée au public. |
| Notes et CRM | interactions, tâches, demandes, historique relationnel, commentaires. | Interne ; zone principale de vigilance RGPD. |
| Incidents et audits | déclarant, email, géolocalisation, photos, observations, scores. | Interne ; peut contenir données incidentelles. |
| Médias | photos, vidéos, crédits, descriptions, analyse IA, droits d'usage. | Peut contenir images de personnes ou métadonnées. |
| Juridique et conformité | SIRET, licences, assurances, justificatifs, dates de validité. | Données pro et pièces potentiellement confidentielles. |
| Logs et versions | auteur, horodatage, avant/après, identifiants techniques. | Preuve de conformité et sécurité. |

## 5. Règles de saisie et publication

Les utilisateurs doivent :

- choisir le champ prévu plutôt que détourner une note libre ;
- séparer strictement les informations publiques, partenaires et internes ;
- contrôler les médias avant publication ;
- ne pas publier les coordonnées personnelles d'un acteur sans base valable et information appropriée ;
- éviter tout commentaire subjectif ou sensible dans les notes ;
- signaler les erreurs de publication contenant des données personnelles ;
- utiliser la modération lorsque le changement a un impact public ou sensible.

## 6. Accès, rôles et permissions

Les droits sont attribués selon l'organisation active, les rôles métier, les rôles admin et les permissions. Le principe de moindre privilège est obligatoire.

Contrôles minimaux :

1. Compte nominatif pour chaque utilisateur.
2. Désactivation rapide au départ d'une personne.
3. Revue trimestrielle des droits admin et permissions sensibles.
4. Revue semestrielle des accès métier.
5. Justification des accès CRM, audit, incident, juridique et logs.
6. Interdiction d'utiliser un compte partagé ou un secret `service_role` côté navigateur.

## 7. Bases légales candidates à confirmer

| Traitement | Base candidate | Validation requise |
|---|---|---|
| Comptes et sécurité | Exécution d'une mission / contrat de travail ou intérêt légitime de sécurité. | Responsable de traitement. |
| Référentiel touristique public | Mission d'intérêt public ou intérêt légitime de promotion touristique selon nature de l'organisme. | Juridique/DPO. |
| Contacts prestataires | Intérêt légitime, contrat, mission publique ou consentement selon contexte. | DPO. |
| Consentements acteurs | Consentement lorsque requis ; preuve dans `actor_consent`. | DPO. |
| CRM relationnel | Intérêt légitime ou mission publique, limité au suivi métier entre l'office et les prestataires. | DPO. |
| Audits/incidents | Mission publique, intérêt légitime, obligation de sécurité selon cas. | DPO/juridique. |
| Logs et sécurité | Obligation de sécurité et intérêt légitime. | DPO. |

## 8. Information des personnes

Chaque personne concernée doit disposer d'une information claire : identité du responsable, finalités, bases légales, données, destinataires, durées, droits, DPO/contact, transferts éventuels et voies de réclamation.

À mettre en place :

- mention d'information pour utilisateurs internes ;
- mention d'information pour prestataires/acteurs ;
- mention pour déclarants d'incident lorsque leur identité est collectée ;
- mention ou clause pour import de données historiques ;
- notice publique sur les fiches ou formulaires concernés.

## 9. Droits des personnes

Le canal officiel de demande doit être publié ou communiqué. Toute demande est transmise au DPO/référent sans délai.

Règles internes :

1. Accuser réception.
2. Vérifier l'identité lorsque nécessaire.
3. Identifier les tables concernées : `auth.users`, `app_user_profile`, `actor`, `actor_channel`, `actor_consent`, `crm_interaction`, `crm_task`, `incident_report`, logs selon limites légales.
4. Répondre dans le délai légal d'un mois, avec prolongation possible uniquement si justifiée et informée.
5. Documenter la réponse et les actions.

## 10. Violations de données

Tout utilisateur doit signaler immédiatement :

- accès non autorisé ;
- publication accidentelle de coordonnées privées ;
- export envoyé au mauvais destinataire ;
- fuite de secret API ;
- perte de données ou altération ;
- média exposant une personne non concernée ;
- erreur RLS ou permission permettant une lecture hors périmètre.

L'équipe DPO/sécurité évalue le risque. Si la violation présente un risque pour les droits et libertés, la notification à la CNIL doit être préparée dans les 72 heures à compter de la connaissance de la violation. Les personnes concernées sont informées si le risque est élevé.

## 11. Conservation et suppression

| Famille | Durée proposée à valider | Commentaire |
|---|---|---|
| Comptes désactivés | À compléter, par exemple durée courte après départ puis anonymisation. | Conserver les traces strictement nécessaires. |
| Contacts prestataires actifs | Tant que la relation est active + revue périodique. | Vérifier l'exactitude. |
| Consentements | Durée de la relation + preuve nécessaire. | Garder l'historique de preuve. |
| CRM et notes privées | À compléter, par exemple 3 à 5 ans selon besoin métier. | Zone prioritaire de minimisation et de purge. |
| Incidents | À compléter selon gravité et obligations. | Ne pas conserver l'identité du déclarant si inutile. |
| Logs d'audit | À compléter selon sécurité/preuve, avec accès restreint. | Tables partitionnées mensuellement. |
| Médias | Tant que les droits sont valides et la diffusion utile. | Surveiller `rights_expires_at`. |
| Pièces légales | Selon durée de validité + prescription applicable. | Ne pas surconserver les justificatifs. |

## 12. Sous-traitants et transferts

La liste des sous-traitants doit être tenue à jour : Supabase/PostgreSQL, OVH France ou autre hébergeur effectif, stockage média, fournisseur d'authentification Google, email/invitation, supervision, sauvegardes et tout service d'analyse média. Le dossier badgeuse indique un hébergement OVH France, dans l'Union européenne, sans transfert hors UE ; ce point doit être confirmé ou adapté pour Bertel.

Pour chaque sous-traitant :

- contrat ou DPA ;
- localisation des données ;
- mesures de sécurité ;
- sous-traitants ultérieurs ;
- mécanisme de transfert hors UE si applicable ;
- procédure de notification d'incident ;
- durée de conservation et suppression.

## 13. Privacy by design dans Bertel

Tout changement produit, DB ou API doit passer une revue courte :

- nouvelles données personnelles ?
- nouveau champ libre ?
- changement de visibilité public/interne ?
- nouveau rôle, permission ou RPC `SECURITY DEFINER` ?
- nouvel export ou partage partenaire ?
- nouveau média ou analyse IA ?
- nouvelle donnée de géolocalisation ?
- impact sur le registre ou l'AIPD ?

## 14. Indicateurs DPO

| Indicateur | Fréquence | Source |
|---|---|---|
| Demandes de droits reçues et délais de réponse | Mensuelle / annuelle | Registre DPO. |
| Violations et incidents de sécurité | À chaque incident + bilan annuel | Registre violations. |
| Nombre d'utilisateurs admin et droits sensibles | Trimestrielle | Page Équipe / RBAC. |
| Fiches avec contacts privés publiés par erreur | Mensuelle | Contrôle qualité. |
| Médias avec droits expirés | Mensuelle | `media.rights_expires_at`. |
| AIPD revues ou mises à jour | À chaque changement majeur | Dossier AIPD. |

{OFFICIAL_SOURCES}

{PROJECT_SOURCES}

{BADGEUSE_SOURCES}
""",
    ),
    PackDoc(
        filename="03_Registre_traitements_RGPD_Bertel",
        title="Registre des traitements RGPD - Bertel",
        subtitle="Fiches de traitements à valider au titre de l'article 30 du RGPD",
        body=f"""\
> Statut : registre initial dérivé du code, du schéma, de la documentation projet et du dossier badgeuse OTI SUD.
> Version : {DATE}. Les informations organisationnelles sont préremplies ; durées, bases légales et sous-traitants effectifs restent à valider.

## 1. Mode d'emploi

Ce registre recense les traitements identifiés dans Bertel. Il doit être conservé par le responsable de traitement ou son DPO, mis à jour après chaque évolution significative et rapproché des politiques de conservation, contrats de sous-traitance, AIPD et procédures de sécurité.

## 2. Informations communes

| Élément | Valeur |
|---|---|
| Responsable de traitement | SPL OTI DU SUD - 379 Rue Hubert Delisle, 97430 Le Tampon - SIREN 882 699 556. |
| DPO / point de contact | David Philippe - Manager SI - d.philippe@otisud.com - 06 93 41 92 91. |
| Hébergement principal | Supabase/PostgreSQL et hébergement de référence OVH France ; architecture Bertel à confirmer si différente. |
| Zone géographique cible | La Réunion / France / UE ; aucun transfert hors UE identifié dans le dossier badgeuse. |
| Personnes concernées | Utilisateurs internes, prestataires touristiques, contacts d'organisations, déclarants d'incident, administrateurs, contributeurs. |
| Mesures transverses | Authentification, RBAC, RLS PostgreSQL, audit log, versioning, modération, séparation public/interne, sauvegardes et supervision à documenter. |

{PUBLIC_DATA_POSTURE}

## 3. Traitement T1 - Authentification, session et profil utilisateur

| Rubrique | Description |
|---|---|
| Finalité | Permettre la connexion, l'identification, la personnalisation et le contrôle d'accès. |
| Données | Email, identifiant utilisateur, nom d'affichage, avatar, locale, fuseau horaire, rôle, préférences, session, MFA éventuel. |
| Tables / services | `auth.users`, `auth.sessions`, `auth.identities`, `app_user_profile`, Supabase Auth, Google OAuth si activé. |
| Personnes | Agents, administrateurs, utilisateurs autorisés. |
| Base légale candidate | Exécution de la mission ou contrat utilisateur ; intérêt légitime de sécurité ; à valider. |
| Accès | Utilisateur concerné, admins habilités, service technique. |
| Conservation | À définir : compte actif, puis durée courte ou archivage limité après départ. |
| Risques | Compte usurpé, droit excessif, maintien de compte inactif. |
| Mesures | Comptes nominatifs, désactivation, revue périodique des droits, politique mot de passe/OAuth, logs. |

## 4. Traitement T2 - Gestion des organisations, rôles et permissions

| Rubrique | Description |
|---|---|
| Finalité | Gérer les équipes, rattachements ORG, rôles métier/admin et permissions. |
| Données | Identifiants utilisateurs, email, organisation active, rôles, rang admin, permissions, statut actif. |
| Tables / services | `user_org_membership`, `user_org_business_role`, `user_org_admin_role`, `org_permission`, `user_permission`, RPC RBAC. |
| Personnes | Membres des offices/organisations et administrateurs. |
| Base légale candidate | Gestion interne et sécurité ; à valider. |
| Accès | Administrateurs habilités, membres pour certaines informations de périmètre. |
| Conservation | À définir : durée d'appartenance + preuve limitée. |
| Risques | Attribution excessive, action sur son propre rôle, accès inter-organisation. |
| Mesures | Rangs admin, anti-self-action, RLS, RPC contrôlées, audit log. |

## 5. Traitement T3 - Référentiel touristique et publication des fiches

| Rubrique | Description |
|---|---|
| Finalité | Créer, maintenir, publier et rechercher des fiches touristiques. |
| Données | Noms d'établissements, descriptions, localisation, horaires, labels, tarifs, offres, classifications, statuts, versions ; données très majoritairement publiques ou publiables. |
| Tables / services | `object`, `object_description`, `object_location`, `object_classification`, `object_taxonomy`, `opening_*`, `object_price`, `publication*`. |
| Personnes | Prestataires lorsqu'ils sont identifiables via une fiche ; utilisateurs auteurs/modérateurs. |
| Base légale candidate | Mission d'intérêt public ou intérêt légitime de promotion touristique ; à valider. |
| Accès | Public pour les données publiées, éditeurs habilités pour brouillons et données étendues. |
| Conservation | Tant que la fiche est utile + archive limitée ; à définir. |
| Risques | Données inexactes, publication d'information interne dans un champ public, rattachement erroné. |
| Mesures | Statuts, modération, versioning, RLS, différenciation public/interne. |

## 6. Traitement T4 - Contacts publics des établissements

| Rubrique | Description |
|---|---|
| Finalité | Diffuser les coordonnées publiques utiles aux visiteurs et partenaires. |
| Données | Téléphone, email, site, réseaux sociaux, rôle du canal, caractère public, primaire. |
| Tables / services | `contact_channel`, `ref_code_contact_kind`, `ref_contact_role`. |
| Personnes | Prestataires, contacts professionnels, établissements. |
| Base légale candidate | Mission publique, intérêt légitime, contrat ou consentement selon type de contact ; à valider. |
| Accès | Public si `is_public`; éditeurs habilités pour gestion complète. |
| Conservation | Tant que le contact est exact et utile ; revue périodique. |
| Risques | Publication d'une coordonnée personnelle ou non destinée au public, notamment si elle diffère des coordonnées professionnelles publiables. |
| Mesures | Champ `is_public`, contrôle éditorial, séparation contact public/contact interne, suppression/correction sur demande. |

## 7. Traitement T5 - Acteurs, contacts privés et consentements

| Rubrique | Description |
|---|---|
| Finalité | Relier des personnes/acteurs à des objets ou organisations et gérer leurs canaux et consentements. |
| Données | Nom, prénom, nom d'affichage, canaux de contact, rôle, organisation, consentement, preuve, préférences. |
| Tables / services | `actor`, `actor_channel`, `actor_object_role`, `actor_consent`. |
| Personnes | Prestataires, gestionnaires, guides, contacts d'organisations. |
| Base légale candidate | Relation professionnelle, intérêt légitime, mission publique ou consentement lorsque requis. |
| Accès | Acteur concerné pour certaines données, utilisateurs habilités selon objet/organisation, admins. |
| Conservation | Tant que la relation est active + durée de preuve consentement ; à définir. |
| Risques | Contact privé exposé, consentement mal qualifié, rattachement erroné. Zone sensible lorsque les coordonnées personnelles diffèrent des coordonnées publiques. |
| Mesures | RLS par objet/acteur, consentement dédié, séparation contact public/acteur. |

## 8. Traitement T6 - Suivi relationnel CRM, notes et relances

| Rubrique | Description |
|---|---|
| Finalité | Historiser les échanges métier entre l'office de tourisme et les prestataires, demandes et relances. |
| Données | Interactions, canal, direction, sujet, notes, dates, propriétaire, acteur, tâches, priorité, statut. |
| Tables / services | `crm_interaction`, `crm_task`, `object_private_description`. |
| Personnes | Contacts prestataires, agents, responsables de suivi. |
| Base légale candidate | Intérêt légitime ou mission publique de suivi relationnel ; à valider. |
| Accès | Interne habilité ; CRM admin/permissions selon implémentation. |
| Conservation | À définir, par exemple 3 à 5 ans selon besoin métier. |
| Risques | Notes subjectives, données sensibles non nécessaires, surconservation ; zone principale de vigilance car elle documente des interactions internes non publiques. |
| Mesures | Règles de rédaction, accès restreint, revue/archivage, audit. |

## 9. Traitement T7 - Modération, versioning et audit log

| Rubrique | Description |
|---|---|
| Finalité | Contrôler les modifications, prouver l'historique, détecter les erreurs et assurer la qualité. |
| Données | Auteur, avant/après, payload, date, statut, revue, snapshots, table modifiée, identifiants. |
| Tables / services | `pending_change`, `object_version`, `audit.audit_log`, triggers d'audit. |
| Personnes | Utilisateurs qui créent, modifient, approuvent ou rejettent. |
| Base légale candidate | Obligation de sécurité, preuve et gouvernance ; à valider. |
| Accès | Modérateurs, admins, service technique. |
| Conservation | À définir selon preuve/sécurité ; accès fortement limité. |
| Risques | Logs trop détaillés contenant des données personnelles, accès admin trop large. |
| Mesures | Partitions, accès admin restreint, minimisation des champs libres, purge à définir. |

## 10. Traitement T8 - Médiathèque et droits d'usage

| Rubrique | Description |
|---|---|
| Finalité | Stocker, qualifier et publier photos, vidéos, documents et visuels touristiques. |
| Données | URL, titre, description, crédit, dimensions, tags, visibilité, droits, analyse IA, portée objet/sous-lieu. |
| Tables / services | `media`, `media_tag`, bucket média, API d'upload, traitement image/vidéo. |
| Personnes | Photographes/crédits, personnes visibles incidentellement, utilisateurs uploaders. |
| Base légale candidate | Publication touristique, contrat/droits d'auteur, intérêt légitime ; à valider. |
| Accès | Public si publié, partenaires/interne selon visibilité, éditeurs habilités. |
| Conservation | Tant que droits valides et diffusion utile ; surveiller expiration. |
| Risques | Image de personne non autorisée, document administratif exposé, droits expirés. |
| Mesures | `is_published`, `visibility`, `rights_expires_at`, revue média, alt text. |

## 11. Traitement T9 - Audits qualité et incidents

| Rubrique | Description |
|---|---|
| Finalité | Réaliser des contrôles qualité et traiter des signalements terrain. |
| Données | Auditeur, critères, scores, notes, incident, gravité, statut, géolocalisation, description, déclarant, email, médias. |
| Tables / services | `audit_template`, `audit_criteria`, `audit_session`, `audit_result`, `incident_report`. |
| Personnes | Auditeurs, déclarants, contacts prestataires, personnes incidentellement présentes dans médias. |
| Base légale candidate | Mission publique, sécurité, intérêt légitime ; à valider. |
| Accès | Interne habilité, admins, agents terrain. |
| Conservation | À définir selon gravité, preuve et suivi opérationnel. |
| Risques | Collecte excessive du déclarant, géolocalisation inutile, photo identifiante. |
| Mesures | Champs facultatifs pour déclarant, limitation média, procédure incident, accès restreint. |

## 12. Traitement T10 - Juridique et conformité prestataires

| Rubrique | Description |
|---|---|
| Finalité | Suivre SIRET, licences, assurances, pièces et statuts de conformité. |
| Données | Type légal, valeur JSON, document, validité, statut, dates de demande/livraison, note. |
| Tables / services | `object_legal`, `ref_legal_type`, `ref_document`, RPC d'audit conformité. |
| Personnes | Prestataires et établissements, parfois dirigeants ou contacts selon document. |
| Base légale candidate | Obligation légale, mission publique, intérêt légitime de conformité ; à valider. |
| Accès | Interne autorisé et B2B restreint selon type ; public seulement si défini comme public. |
| Conservation | Selon validité et prescriptions applicables. |
| Risques | Pièces trop sensibles, conservation après expiration, diffusion non autorisée. |
| Mesures | Typage public/interne, statut, dates, revue d'expiration, RLS. |

## 13. Traitement T11 - Imports, migration et rapprochement de données historiques

| Rubrique | Description |
|---|---|
| Finalité | Importer, contrôler et promouvoir des données historiques ou candidates vers le référentiel. |
| Données | Données de fiches, contacts, médias, lots, événements d'import, décisions de mapping, erreurs. |
| Tables / services | `staging.*`, `import_batches`, `import_events`, `mapping_*`, RPC `commit_staging_to_public`. |
| Personnes | Prestataires et contacts présents dans sources historiques, agents de migration. |
| Base légale candidate | Continuité du référentiel / mission publique ; à valider. |
| Accès | Équipe data, admins et métiers habilités. |
| Conservation | Données staging à purger après validation ; durée à fixer. |
| Risques | Import d'informations obsolètes, doublons, absence d'information des personnes. |
| Mesures | Profilage, validation, ledger, règles de promotion, purge des lots temporaires. |

## 14. Traitement T12 - Tableaux de bord, recherche et reporting

| Rubrique | Description |
|---|---|
| Finalité | Mesurer qualité, actualisation, couverture, activité et état du référentiel. |
| Données | Agrégats, filtres, catégories, dates de mise à jour, indicateurs qualité, objets visibles. |
| Tables / services | `internal.mv_filtered_objects`, RPC dashboard, caches et vues matérialisées. |
| Personnes | Utilisateurs internes, prestataires indirectement via fiches agrégées. |
| Base légale candidate | Pilotage métier et qualité ; à valider. |
| Accès | Utilisateurs habilités selon tableau de bord. |
| Conservation | Agrégats recalculés ; pas de conservation séparée sauf exports. |
| Risques | Export non maîtrisé ou ré-identification via filtres fins. |
| Mesures | RLS, limitation des exports, partage contrôlé. |

## 15. Plan d'action registre

| Priorité | Action | Responsable |
|---|---|---|
| Haute | Confirmer que le responsable de traitement et le référent RGPD repris du dossier badgeuse s'appliquent à Bertel. | Direction / DPO |
| Haute | Valider les bases légales par traitement, en distinguant données publiques et zones internes CRM/contacts privés. | DPO / juridique |
| Haute | Fixer les durées de conservation et modalités de purge/anonymisation, surtout pour CRM, notes, imports et logs. | DPO / produit / technique |
| Moyenne | Finaliser la liste des sous-traitants et transferts. | DPO / technique |
| Moyenne | Créer un registre des demandes de droits et un registre des violations. | DPO |
| Moyenne | Mettre en place la revue périodique des accès. | Admin plateforme |

{OFFICIAL_SOURCES}

{PROJECT_SOURCES}

{BADGEUSE_SOURCES}
""",
    ),
    PackDoc(
        filename="04_AIPD_Bertel",
        title="Analyse d'impact relative à la protection des données - Bertel",
        subtitle="AIPD de précaution centrée sur la distinction données publiques / données internes",
        body=f"""\
> Statut : AIPD initiale pré-remplie à partir du code et du schéma. Elle doit être revue avec les responsables métier, sécurité et DPO avant validation.
> Version : {DATE}. Revue recommandée à chaque changement majeur de données, accès, média, IA, export ou sous-traitant.

## 1. Décision de réaliser une AIPD

L'AIPD est maintenue comme analyse de précaution et de documentation pour Bertel. Le traitement n'apparaît pas, à ce stade, comme principalement composé de données à haut risque : environ 95 % des informations stockées sont publiques ou destinées à être publiées dans un référentiel touristique. Les points de vigilance réels se concentrent sur les interactions internes entre l'office de tourisme et les prestataires, les notes CRM, les tâches, les historiques de relation, les coordonnées personnelles lorsqu'elles diffèrent des coordonnées publiques ou professionnelles, ainsi que les erreurs de publication accidentelles.

Conclusion provisoire : AIPD utile à maintenir pour démontrer la proportionnalité et piloter les mesures. Aucun risque élevé résiduel n'est identifié à ce stade si les mesures prévues sont appliquées ; la consultation préalable de la CNIL n'est donc pas anticipée, sauf évolution du périmètre ou impossibilité de réduire un risque élevé.

{PUBLIC_DATA_POSTURE}

## 2. Traitement étudié

| Élément | Description |
|---|---|
| Nom | Bertel - référentiel touristique, publication, relation prestataires et gouvernance qualité. |
| Responsable de traitement | SPL OTI DU SUD - 379 Rue Hubert Delisle, 97430 Le Tampon - SIREN 882 699 556. |
| DPO | David Philippe - Manager SI - d.philippe@otisud.com - 06 93 41 92 91. |
| Finalités | Référentiel touristique, publication, qualité des données, relation prestataires, conformité, sécurité, modération et reporting. |
| Personnes concernées | Utilisateurs internes, prestataires, contacts d'organisations, déclarants d'incident, personnes visibles dans médias. |
| Données principales | Données publiques ou publiables majoritaires : fiches touristiques, coordonnées professionnelles d'établissement, horaires, tarifs, descriptions, médias de promotion, classifications. Données internes minoritaires : profils utilisateurs, rôles, notes, tâches, interactions CRM, consentements, incidents, logs, versions, pièces légales. |
| Destinataires | Public pour données publiées, utilisateurs habilités, administrateurs, partenaires autorisés, sous-traitants techniques. |
| Technologies | Next.js, Supabase Auth, PostgreSQL/RLS, PostGIS, stockage média, API/RPC, imports staging, traitement images/vidéos ; hébergement de référence OVH France à confirmer pour Bertel. |

## 3. Qualification du caractère public des données

Le point structurant de l'AIPD Bertel est que le traitement porte principalement sur un référentiel touristique public ou publiable. Les données publiques ne disparaissent pas du champ RGPD lorsqu'elles identifient une personne, mais leur niveau de risque est nettement plus faible lorsqu'elles correspondent à des coordonnées professionnelles ou à des informations volontairement diffusées pour la promotion touristique. L'analyse d'impact doit donc éviter de qualifier l'ensemble de Bertel comme un traitement sensible : elle doit concentrer les mesures sur la frontière entre données publiques et données internes.

<!-- PAGE_BREAK -->

| Famille de données | Statut dans Bertel | Impact AIPD |
|---|---|---|
| Fiches touristiques, noms d'établissements, descriptions, horaires, tarifs, équipements, labels, classifications. | Public ou destiné à publication. | Risque RGPD intrinsèque faible ; enjeux principaux : exactitude, mise à jour, cohérence de publication. |
| Coordonnées professionnelles d'établissement : téléphone, email générique, site web, réseaux sociaux officiels. | Public si le canal est officiel et destiné aux visiteurs. | Risque faible si la nature professionnelle est confirmée ; risque moyen si une coordonnée personnelle est publiée par erreur. |
| Médias de promotion touristique, crédits et droits d'usage. | Publiable selon droits, crédit, visibilité et expiration. | Risque principalement lié au droit à l'image, aux personnes incidentelles et aux documents visibles. |
| Interactions entre l'office de tourisme et les prestataires. | Interne, non destiné à publication. | Zone de vigilance principale : contexte relationnel, historique, demandes, relances, appréciations éventuelles. |
| Notes CRM, tâches, commentaires, descriptions privées. | Interne, accès restreint. | Zone de vigilance principale : subjectivité, données excessives, conservation trop longue. |
| Coordonnées personnelles différentes des coordonnées publiques. | Interne sauf base légale ou consentement spécifique. | Risque principal d'exposition accidentelle ou d'usage non prévu. |
| Comptes utilisateurs, rôles, permissions, logs, versions. | Interne technique et organisationnel. | Risque de sécurité, accès excessif, traçabilité et conservation. |

Conséquence pour l'AIPD : le risque général du référentiel public est faible. Les mesures prioritaires portent sur la distinction visible entre champs publics et champs internes, la validation des coordonnées personnelles, la discipline de rédaction des notes CRM, les exports et la conservation.

<!-- PAGE_BREAK -->

## 4. Flux de données

1. Un utilisateur se connecte via Supabase Auth ou Google OAuth.
2. La session hydrate le rôle plateforme, l'organisation active, le rang admin et les permissions.
3. Les fiches touristiques sont consultées via Explorer, dashboard ou API.
4. Les éditeurs créent ou modifient des fiches ; certains changements passent par `pending_change`.
5. Les médias sont uploadés, traités, rattachés à un objet ou sous-lieu et publiés selon droits.
6. Les acteurs, contacts, consentements, notes et relances alimentent le suivi relationnel.
7. Les audits et incidents peuvent créer des tâches ou interactions.
8. Les triggers et logs conservent versions, avant/après et traces d'administration.
9. Les imports staging rapprochent des données historiques avant promotion en production.

## 5. Nécessité et proportionnalité

| Question | Évaluation provisoire |
|---|---|
| Les données sont-elles nécessaires ? | Oui. La majorité des données correspond au référentiel public ou publiable nécessaire à la mission de promotion touristique. Les données internes CRM/contacts privés doivent rester limitées au suivi métier réel. |
| Les finalités sont-elles explicites ? | Oui : publication touristique, qualité du référentiel, relation prestataires, conformité et sécurité. Elles doivent être formalisées dans les mentions d'information. |
| Les bases légales sont-elles documentées ? | Partiellement ; validation DPO requise par traitement. |
| Les personnes sont-elles informées ? | À compléter : mentions utilisateurs, prestataires, acteurs et déclarants d'incident ; préciser clairement que les informations publiques/professionnelles peuvent être diffusées dans le référentiel. |
| Les droits sont-ils exerçables ? | Procédure à mettre en place ; tables cibles identifiées. |
| La conservation est-elle limitée ? | À compléter : durées et purge/anonymisation non encore formalisées, priorité sur CRM, notes privées, imports staging et logs. |
| Les accès sont-ils maîtrisés ? | Oui techniquement par rôles, RBAC et RLS ; revue périodique à instituer. |
| Les données publiques/interne sont-elles séparées ? | Oui par champs `is_public`, `visibility`, statuts et politiques ; vigilance sur notes/médias. |

## 6. Mesures existantes identifiées

- Authentification Supabase Auth et option Google OAuth.
- Communications HTTPS.
- Rôles plateforme : `owner`, `super_admin`, `tourism_agent`.
- Rôles organisation, rangs admin et permissions fines.
- RLS sur les tables publiques exposées et helpers `api.can_read_object`, `api.user_can_write_object_canonical`, `api.current_user_extended_object_ids`.
- Séparation publié/brouillon/archivé/masqué.
- Modération via `pending_change`.
- Versioning `object_version`.
- Audit log partitionné `audit.audit_log` pour UPDATE/DELETE.
- Logs administratifs et contrôles d'accès, selon le modèle badgeuse.
- Triggers de cohérence : objet, média principal, dates légales, statut, ORG.
- Visibilité médias et contacts.
- Consentements dédiés pour acteurs.
- Indicateurs de droits médias (`rights_expires_at`).
- Hébergement UE de référence OVH France, à confirmer si l'architecture Bertel diffère.

## 7. Risques et mesures complémentaires

| Risque | Niveau initial | Mesures existantes | Mesures à ajouter | Résiduel cible |
|---|---|---|---|---|
| Stockage et publication du référentiel touristique public. | Faible | Statuts de publication, champs de visibilité, modération, RLS, versioning. | Contrôle qualité périodique et mentions d'information indiquant la diffusion des données publiques/professionnelles. | Faible |
| Publication d'une coordonnée personnelle non destinée au public. | Moyen | `is_public`, séparation `contact_channel` / `actor_channel`, RLS. | Revue périodique contacts, avertissement UI, procédure correction urgente, règle de distinction coordonnées publiques/personnelles. | Faible à moyen |
| Accès non autorisé aux interactions OTI/prestataires, brouillons ou notes privées. | Moyen | RLS, organisation active, `can_read_extended`, RBAC. | Tests RLS périodiques, revue admin, journal des exceptions. | Faible |
| Notes CRM contenant données sensibles, subjectives ou non nécessaires. | Moyen | Accès interne, page CRM, permissions. | Charte de rédaction, purge/archivage, formation utilisateurs. | Faible à moyen |
| Confusion entre donnée publique et donnée interne lors d'un export ou partage. | Moyen | Statuts, champs de visibilité, séparation public/interne. | Checklist export, revue avant diffusion, limitation des exports massifs. | Faible |
| Secret technique exposé dans le frontend ou dépôt. | Moyen | Séparation client/service role prévue. | Scan secrets, rotation, revue CI/CD, interdiction `service_role` navigateur. | Faible |
| Média publiant une personne ou document identifiable sans droit. | Moyen | `is_published`, `visibility`, droits, crédits. | Checklist média, floutage, validation avant publication, expiration automatisée. | Faible à moyen |
| Incident géolocalisé ou photo identifiant un déclarant ou tiers. | Faible à moyen | Champs déclarant facultatifs, accès interne. | Formulaire minimisé, revue média incident, suppression des identifiants inutiles. | Faible |
| Conservation excessive des comptes, notes, logs ou imports staging. | Moyen | Partitions logs, statuts, archivage notes. | Politique de conservation, jobs de purge/anonymisation, revue DPO. | Faible à moyen |
| Données historiques importées obsolètes ou mal qualifiées public/interne. | Faible à moyen | Staging, mapping, ledger de promotion. | Notice import, contrôle exactitude, qualification public/interne, purge des lots rejetés. | Faible |
| Analyse IA des médias produisant des inférences non prévues. | Faible à moyen | `analyse_data` isolé, usage décrit comme classification. | Encadrer finalité IA, limiter conservation, informer si usage significatif. | Faible |
| Indisponibilité ou perte de données. | Moyen | Supabase/PostgreSQL, sauvegardes à confirmer. | Plan sauvegarde-restauration, tests de restauration, RTO/RPO. | Faible |

## 8. Évaluation des droits et libertés

Les impacts possibles portent sur :

- exposition non souhaitée de coordonnées personnelles non destinées au public ;
- perte de maîtrise sur les contacts et consentements ;
- appréciations inexactes ou subjectives dans les notes CRM ;
- refus ou retard de correction/suppression ;
- réutilisation de médias sans droit ;
- accès non autorisé par un membre hors périmètre ;
- conservation trop longue de traces et données historiques.

Les catégories de données ne visent pas des données sensibles au sens de l'article 9 du RGPD, et le référentiel est très majoritairement public ou publiable. Cela réduit fortement le risque général du traitement : l'AIPD ne considère pas la simple conservation de fiches touristiques publiques comme un risque élevé. Des champs libres, interactions internes, coordonnées personnelles non publiques et médias peuvent toutefois contenir accidentellement des données personnelles non nécessaires. Cette possibilité doit être traitée comme un risque opérationnel ciblé. Le niveau résiduel visé est faible à moyen ; aucun risque élevé résiduel n'est identifié à ce stade sous réserve de validation DPO et de mise en oeuvre des mesures.

## 9. Plan d'action AIPD

| Priorité | Action | Responsable proposé |
|---|---|---|
| Haute | Confirmer responsable de traitement, référent RGPD, bases légales et mentions d'information. | DPO / direction |
| Haute | Formaliser les durées de conservation et procédures de purge/anonymisation, surtout CRM, notes, coordonnées personnelles non publiques, imports et logs. | DPO / technique |
| Haute | Créer procédure violation et registre des violations. | DPO / sécurité |
| Moyenne | Auditer les RLS et les RPC `SECURITY DEFINER` sensibles. | Technique |
| Moyenne | Ajouter une checklist public/interne pour publication média/contact dans l'UI ou la procédure. | Produit |
| Moyenne | Former les utilisateurs aux notes CRM factuelles et minimisées. | DPO / métier |
| Moyenne | Documenter sous-traitants, transferts, sauvegardes et restauration ; confirmer OVH France ou l'hébergement effectif Bertel. | Technique / DPO |
| Moyenne | Définir revue trimestrielle des droits admin et permissions. | Admin plateforme |

## 10. Critères de révision

Réviser l'AIPD lorsque :

- nouvelle famille de données personnelles ;
- nouveau champ libre, média ou analyse IA ;
- nouvelle diffusion publique ou export partenaire ;
- nouveau sous-traitant ou transfert hors UE ;
- changement significatif de RLS/RBAC/RPC ;
- incident ou violation révélant un risque non couvert ;
- extension à de nouveaux publics ou organisations.

## 11. Avis et validation

| Étape | Statut | Commentaire |
|---|---|---|
| Avis métier | À compléter | Validation des finalités et usages. |
| Avis technique/sécurité | À compléter | Validation mesures d'accès, logs, sauvegardes. |
| Avis DPO | À compléter | Confirmer AIPD de précaution, risques résiduels faibles à moyens, plan d'action. |
| Consultation CNIL préalable | Non envisagée à ce stade | Requise seulement si risque élevé résiduel impossible à réduire. |
| Date de prochaine revue | À compléter | Recommandé : au moins annuelle ou changement majeur. |

{OFFICIAL_SOURCES}

{PROJECT_SOURCES}

{BADGEUSE_SOURCES}
""",
    ),
    PackDoc(
        filename="05_Procedures_DPO_Bertel",
        title="Procédures DPO opérationnelles - Bertel",
        subtitle="Demandes de droits, violations, registre, AIPD, accès, conservation et contrôles",
        body=f"""\
> Statut : procédures de travail à adapter à l'organisation. Les délais légaux doivent être pilotés par le DPO ou le référent désigné.
> Version : {DATE}.

{ORG_CONTEXT}

{PUBLIC_DATA_POSTURE}

## 1. Procédure de tenue du registre

Déclencheurs de mise à jour :

- nouvelle table ou nouveau champ personnel ;
- nouveau module UI ou API ;
- nouveau sous-traitant ;
- nouvel export ou partage ;
- changement de base légale, durée ou destinataire ;
- incident révélant un traitement non documenté.

Étapes :

1. Identifier le traitement concerné.
2. Mettre à jour la fiche registre.
3. Vérifier les mentions d'information.
4. Évaluer si l'AIPD doit être créée ou revue.
5. Faire valider par le DPO.
6. Conserver la preuve de validation.

## 2. Procédure de demande de droits

Canaux possibles : adresse DPO, formulaire, courrier, demande reçue par un agent, demande via prestataire.

1. Enregistrer la demande : date, identité, canal, droit exercé, périmètre.
2. Accuser réception.
3. Vérifier l'identité uniquement si nécessaire et de manière proportionnée.
4. Identifier les tables : `auth.users`, `app_user_profile`, `actor`, `actor_channel`, `actor_consent`, `contact_channel`, `crm_interaction`, `crm_task`, `incident_report`, `object_private_description`, logs selon limites.
5. Geler toute suppression automatique qui pourrait empêcher la réponse.
6. Préparer la réponse : accès, rectification, effacement, limitation, opposition, portabilité si applicable.
7. Répondre dans le délai d'un mois. Si la demande est complexe, documenter la prolongation et informer la personne dans le délai initial.
8. Journaliser la réponse et les actions effectuées.

| Champ de registre droits | Exemple |
|---|---|
| Référence | DDP-2026-001 |
| Date de réception | À compléter |
| Demandeur | À compléter |
| Droit exercé | Accès / rectification / effacement / opposition / limitation |
| Tables concernées | À compléter |
| Responsable interne | DPO |
| Échéance un mois | À compléter |
| Réponse envoyée | À compléter |
| Action technique | À compléter |

## 3. Procédure de violation de données personnelles

Une violation peut être une destruction, perte, altération, divulgation ou accès non autorisé à des données personnelles.

Déclencheurs Bertel :

- contact privé publié ;
- interaction OTI/prestataire ou note CRM rendue visible hors périmètre ;
- média exposant une personne ou document personnel ;
- erreur RLS donnant accès à des données internes ;
- invitation ou rôle attribué à la mauvaise personne ;
- export envoyé au mauvais destinataire ;
- secret Supabase ou token exposé ;
- perte ou altération de données ;
- accès suspect dans les logs.

Étapes immédiates :

1. Contenir : retirer la publication, couper l'accès, révoquer token, désactiver compte si nécessaire.
2. Préserver les preuves : horodatage, captures internes, logs, personnes impliquées.
3. Qualifier : données, personnes, volume, durée d'exposition, destinataires, contexte.
4. Évaluer le risque pour les droits et libertés.
5. Décider notification CNIL dans les 72 heures si risque.
6. Informer les personnes concernées si risque élevé.
7. Corriger la cause racine et documenter.
8. Clôturer avec mesures préventives.

| Champ registre violation | Valeur |
|---|---|
| Référence | VIOL-YYYY-NNN |
| Date de connaissance | À compléter |
| Description | À compléter |
| Données concernées | À compléter |
| Personnes concernées | À compléter |
| Cause | À compléter |
| Mesures immédiates | À compléter |
| Risque | Aucun / risque / risque élevé |
| CNIL notifiée | Oui / non / motif |
| Personnes informées | Oui / non / motif |
| Actions préventives | À compléter |

## 4. Procédure d'accès, arrivée et départ utilisateur

Arrivée :

1. Valider le besoin métier et l'organisation.
2. Créer ou inviter le compte nominatif.
3. Attribuer le rôle métier minimal.
4. Ajouter un rôle admin seulement si nécessaire.
5. Vérifier les permissions effectives.
6. Informer l'utilisateur du règlement RGPD interne.

Départ ou changement de poste :

1. Désactiver l'appartenance organisation.
2. Révoquer rôles admin et permissions.
3. Vérifier les tâches/objets assignés.
4. Conserver ou anonymiser le profil selon politique de conservation.
5. Documenter l'action.

Revue périodique :

- droits admin : trimestriel ;
- permissions sensibles : trimestriel ;
- membres actifs : semestriel ;
- comptes sans activité : à définir.

## 5. Procédure de publication contact et média

Avant publication d'un contact :

- le canal est-il professionnel ou destiné au public ?
- la coordonnée est-elle bien publique/professionnelle, et non personnelle ou réservée aux échanges internes ?
- `is_public` est-il correct ?
- la personne est-elle informée ?
- le contact est-il exact et à jour ?
- le rôle du canal est-il renseigné ?

Avant publication d'un média :

- droit d'usage documenté ;
- crédit renseigné si nécessaire ;
- absence de personne identifiable non autorisée ;
- absence de plaque, document, email ou numéro privé ;
- visibilité correcte ;
- date d'expiration contrôlée ;
- description accessible si diffusion publique.

## 6. Procédure AIPD et privacy by design

Créer ou réviser l'AIPD si :

- nouveau traitement potentiellement à risque ;
- changement de finalité ;
- nouvelle analyse IA ou enrichissement automatisé ;
- nouvelle géolocalisation ou média sensible ;
- élargissement de diffusion ;
- changement d'hébergement/sous-traitant ;
- incident significatif ;
- nouveau public de personnes concernées.

Checklist de revue :

1. Finalité claire.
2. Base légale candidate.
3. Données minimales.
4. Accès et RLS.
5. Durée de conservation.
6. Information des personnes.
7. Droits exerçables.
8. Sous-traitants et transferts.
9. Risques et mesures.
10. Validation DPO.

## 7. Procédure imports et données historiques

1. Identifier la source et le responsable de la source.
2. Profiler les colonnes personnelles.
3. Supprimer les colonnes inutiles avant import.
4. Charger en staging.
5. Contrôler doublons, exactitude et obsolescence.
6. Documenter mapping et décisions.
7. Promouvoir uniquement les données nécessaires.
8. Purger ou archiver les lots rejetés selon durée validée.
9. Mettre à jour le registre si la finalité change.

## 8. Procédure conservation, purge et anonymisation

La politique de conservation doit être validée avant automatisation. À défaut, aucune purge massive ne doit être lancée sans accord DPO/technique.

Ordre de travail :

1. Définir les durées par famille de données.
2. Définir ce qui doit rester en preuve légale ou sécurité.
3. Choisir suppression, anonymisation ou archivage restreint.
4. Tester sur environnement non production.
5. Vérifier impact sur intégrité, audits et versioning.
6. Journaliser l'opération.
7. Mettre à jour le registre.

<!-- PAGE_BREAK -->

## 9. Modèle de mention d'information courte

À adapter selon canal :

| Élément | Texte à adapter |
|---|---|
| Responsable | Vos données sont traitées par SPL OTI DU SUD, 379 Rue Hubert Delisle, 97430 Le Tampon, pour la gestion du référentiel touristique Bertel. |
| Finalités | Publication des informations touristiques, relation prestataires, qualité, conformité et sécurité. |
| Données | Coordonnées professionnelles, informations de relation, traces d'administration et, le cas échéant, signalements. |
| Base légale | À préciser selon le traitement : mission publique, intérêt légitime, contrat, obligation légale ou consentement. |
| Destinataires | Agents habilités, administrateurs, partenaires autorisés et sous-traitants techniques. Certaines données peuvent être publiées si elles sont destinées au public. |
| Durée | Les données sont conservées selon une politique validée, proportionnée à la finalité. |
| Droits | Vous pouvez exercer vos droits auprès de David Philippe, référent RGPD, d.philippe@otisud.com. Vous pouvez aussi saisir la CNIL. |

## 10. Rapport annuel DPO - indicateurs suggérés

| Indicateur | Donnée attendue |
|---|---|
| Demandes de droits | Nombre, type, délai moyen, retard éventuel. |
| Violations | Nombre, niveau de risque, notification CNIL/personnes, actions préventives. |
| AIPD | Créées, revues, actions ouvertes/fermées. |
| Accès | Nombre d'admins, revues réalisées, comptes désactivés. |
| Données | purges, corrections, contacts/médias revus. |
| Formation | Utilisateurs formés au RGPD et aux notes CRM. |
| Sous-traitants | Revue DPA, transferts, incidents fournisseur. |

{OFFICIAL_SOURCES}

{PROJECT_SOURCES}

{BADGEUSE_SOURCES}
""",
    ),
]


def full_markdown(doc: PackDoc) -> str:
    return f"# {doc.title}\n\n{doc.subtitle}\n\n{doc.body.rstrip()}\n"


def set_cell_width(cell, width_in: float) -> None:
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, width_in: float = 6.5) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_runs(paragraph, text: str) -> None:
    # Minimal Markdown inline support for **bold** and `code`.
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    add_runs(paragraph, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    set_table_width(table)

    if col_count == 2:
        widths = [1.75, 4.75]
    elif col_count == 3:
        widths = [1.55, 2.25, 2.7]
    elif col_count == 4:
        widths = [1.15, 1.75, 2.15, 1.45]
    else:
        widths = [6.5 / col_count] * col_count

    for c_idx in range(col_count):
        cell = table.rows[0].cells[c_idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[c_idx])
        shade_cell(cell, "E8EEF5")
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(rows[0][c_idx].strip())
        run.bold = True

    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    prevent_row_split(table.rows[0])

    for source_row in rows[1:]:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for c_idx in range(col_count):
            text = source_row[c_idx].strip() if c_idx < len(source_row) else ""
            cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cells[c_idx], widths[c_idx])
            cells[c_idx].text = ""
            paragraph = cells[c_idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_runs(paragraph, text)


def add_numbered_text(doc: Document, line: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_runs(paragraph, line)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip()
        parts = [part.strip() for part in line.strip("|").split("|")]
        if not all(set(part) <= {"-", ":", " "} for part in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def configure_styles(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)

    footer = section.footer.paragraphs[0]
    footer.text = f"Bertel - Pack RGPD/DPO - {DATE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)


def markdown_to_docx(markdown: str, out_path: Path, title: str, subtitle: str) -> None:
    doc = Document()
    configure_styles(doc, title)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(75, 85, 99)

    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue
        if line == "<!-- PAGE_BREAK -->":
            doc.add_page_break()
            idx += 1
            continue
        if line.startswith("# "):
            idx += 1
            continue
        if line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), "Heading 1")
        elif line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), "Heading 2")
        elif line.startswith("#### "):
            add_paragraph(doc, line[5:].strip(), "Heading 3")
        elif line.startswith("> "):
            chunks = []
            while idx < len(lines) and lines[idx].strip().startswith("> "):
                chunks.append(lines[idx].strip()[2:].strip())
                idx += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            shade_paragraph(paragraph, "F4F6F9")
            add_runs(paragraph, " ".join(chunks))
            idx -= 1
        elif line.startswith("- "):
            add_paragraph(doc, line[2:].strip(), "List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            add_numbered_text(doc, line)
        elif line.startswith("|"):
            rows, new_idx = parse_table(lines, idx)
            add_table(doc, rows)
            idx = new_idx - 1
        else:
            paragraph_text = line
            while (
                idx + 1 < len(lines)
                and lines[idx + 1].strip()
                and not lines[idx + 1].strip().startswith(("#", "-", ">", "|"))
                and not re.match(r"^\d+\.\s+", lines[idx + 1].strip())
            ):
                idx += 1
                paragraph_text += " " + lines[idx].strip()
            add_paragraph(doc, paragraph_text)
        idx += 1

    doc.save(out_path)


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for doc in DOCUMENTS:
        markdown = full_markdown(doc)
        (ROOT / doc.md_name).write_text(markdown, encoding="utf-8")
        markdown_to_docx(markdown, OUT / doc.docx_name, doc.title, doc.subtitle)
    index = "\n".join(
        [
            "# Pack documentaire RGPD/DPO Bertel",
            "",
            f"Version : {DATE}",
            "",
            "## Livrables",
            "",
            *[f"- `{doc.md_name}` + `livrables/{doc.docx_name}` : {doc.subtitle}" for doc in DOCUMENTS],
            "",
            "## À valider par le DPO",
            "",
            "- confirmation que SPL OTI DU SUD et le référent RGPD repris du dossier badgeuse s'appliquent bien à Bertel ;",
            "- bases légales par traitement ;",
            "- durées de conservation, surtout CRM, notes, imports et logs ;",
            "- sous-traitants, transferts et contrats, notamment confirmation de l'hébergement effectif Bertel ;",
            "- niveau de risque résiduel de l'AIPD, a priori faible à moyen si les mesures sont appliquées ;",
            "- mentions d'information et canal d'exercice des droits.",
            "",
            OFFICIAL_SOURCES.strip(),
            "",
            PROJECT_SOURCES.strip(),
            "",
            BADGEUSE_SOURCES.strip(),
            "",
        ]
    )
    (ROOT / "00_INDEX_PACK_RGPD_DPO.md").write_text(index, encoding="utf-8")
    markdown_to_docx(index, OUT / "00_INDEX_PACK_RGPD_DPO.docx", "Pack documentaire RGPD/DPO Bertel", "Index des documents, sources et points à valider")


if __name__ == "__main__":
    build()
